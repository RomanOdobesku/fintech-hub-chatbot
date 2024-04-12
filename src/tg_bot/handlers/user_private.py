import docx
import tempfile
import datetime
from aiogram import Bot
from aiogram import Router, F
from aiogram.types import FSInputFile
from aiogram.filters import Command
from aiogram.types import Message, CallbackQuery
from aiogram.enums.parse_mode import ParseMode
from sqlalchemy.ext.asyncio import AsyncSession
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.tg_bot.kb import reply
import src.tg_bot.bot_database.db_query as db_query

user_private_router = Router()


@user_private_router.message(Command("start"))
async def start_handler(msg: Message, session: AsyncSession) -> None:
    """
    Обработчик команды "/start". Отправляет приветственное сообщение пользователю
    и добавляет его в базу данных.

    Аргументы:
    - msg (Message): Объект сообщения от пользователя.
    - session (AsyncSession): Сессия асинхронного соединения с базой данных.

    Возвращает:
    - None
    """
    await msg.answer("""Привет 👻, я виртуальный помощник для отслеживания новостей по разными темам в финтех-сфере! \n\nПерейди к выбору категорий по кнопке ниже, отметь подходящие темы и нажми кнопку finish, и я буду присылать тебе самые свежие новости!""",
                     reply_markup=reply.start_kb)
    await db_query.orm_add_user(session, str(msg.from_user.id), str(msg.from_user.first_name))


@user_private_router.message(Command("restart"))
async def restart(msg: Message) -> None:
    """
    Обработчик команды "/restart". Выводит inline-клавиатуру "да/нет".

    Аргументы:
    - msg (Message): Объект сообщения от пользователя.

    Возвращает:
    - None
    """
    await msg.answer("Ты уверен, что хотите сбросить текущие подписки?",
                     reply_markup=reply.restart_categories_kb)


@user_private_router.message(Command("my_categories"))
async def print_categories(msg: Message, session: AsyncSession) -> None:
    """
    Обработчик команды "/my_categories". Выводит сообщение с перечисленными названиями категорий.

    Аргументы:
    - msg (Message): Объект сообщения от пользователя.
    - session (AsyncSession): Сессия асинхронного соединения с базой данных.

    Возвращает:
    - None
    """
    users_id = str(msg.from_user.id)
    cats = await db_query.orm_get_users_categories(session, users_id)
    names = [await db_query.orm_get_name_category_by_id(session, element) for element in cats]
    answer_message = "*Ваш список тем:*\n"
    for element in names:
        answer_message += f"- {element}\n"
    await msg.answer(answer_message, parse_mode=ParseMode.MARKDOWN)


@user_private_router.message(Command("help"))
async def help_handler(msg: Message) -> None:
    """
    Обработчик команды "/help". Выводит справку о боте.

    Аргументы:
    - msg (Message): Объект сообщения от пользователя.

    Возвращает:
    - None
    """
    await msg.answer("""Если ты уже выбрал категории, то можешь получить составленный для тебя дайджест с помощью команды /digest, либо дождаться ежедневной рассылки (в 09:00). \nПосмотреть выбранные категории ты можешь по команде /my_categories, а сбросить прежние настройки и выбрать новые темы с помощью /restart.""")


@user_private_router.message(Command("digest"))
async def send_news_to_subscribers(msg: Message, session: AsyncSession, bot: Bot) -> None:
    """
    Обработчик команды "/digest". Выводит по 5 недавних новостей
    по категориям, выбранных пользователем.

    Аргументы:
    - msg (Message): Объект сообщения от пользователя.
    - session (AsyncSession): Сессия асинхронного соединения с базой данных.

    Возвращает:
    - None
    """
    users_chat_id = await db_query.orm_get_list_of_users(session)

    latest_news = await db_query.orm_get_latest_news_by_categories(session)
    category = await db_query.orm_get_list_of_category(session)
    
    for user in users_chat_id:
        message = ""
        subscriptions = await db_query.orm_get_users_categories(session, str(user))
        for subscription in subscriptions:
            # создаем сообщение для рассылки
            message += "\n*{}* \n".format(category[int(subscription)])
            for news in latest_news[int(subscription)][:5]:
                    message += f"\n - {news.title}. [Читать в источнике]({news.url})\n"
            message += '\n'

        await msg.answer(message, parse_mode=ParseMode.MARKDOWN, disable_web_page_preview=True)

        await send_docx_to_subscribers(session, bot)


@user_private_router.callback_query(F.data.startswith('print_help'))
async def print_help(callback: CallbackQuery):
    """
    Обработчик нажатия на кнопку "Помощь" стартовой клавиатуры. Выводит справку о боте.

    Аргументы:
    - callback (CallbackQuery): Объект сообщения от пользователя.

    Возвращает:
    - None
    """
    await callback.message.edit_reply_markup()
    await callback.message.delete()
    await callback.message.answer("""Выбери интересующие тебя категории с помощью кнопки "меню", и я буду присылать ежедневные дайджесты с важными актуальными новостями по интересующим тебя темам!""")


@user_private_router.callback_query(F.data.startswith('choose_category'))
async def category_hadnler(callback: CallbackQuery) -> None:
    """
    Обработчик нажатия на кнопку "Выбрать категорию". Выводит клавиатуру для выбора категорий.

    Аргументы:
    - callback (CallbackQuery): Объект сообщения от пользователя.
    Возвращает:
    - None
    """
    await callback.message.edit_reply_markup()
    await callback.message.delete()
    await callback.message.answer("Выбери категории для отслеживания:",
                                  reply_markup=reply.categories_kb)


@user_private_router.callback_query(F.data.startswith('subscribe_finish'))
async def subscribe_finish(callback: CallbackQuery,
                           session: AsyncSession,
                           apscheduler: AsyncIOScheduler,
                           bot: Bot):
    """
    Обработчик нажатия на кнопку "finish". Убирает inline-клавиатуру с выбором категорий
    и выводит сообщение о сохранении подписок.

    Аргументы:
    - callback (CallbackQuery): Объект сообщения от пользователя.
    - session (AsyncSession): Сессия асинхронного соединения с базой данных.
    - apscheduler (AsyncIOScheduler): Scheduler для запуска автоматической рассылки дайджеста.
    - bot (Bot): Объект класса Bot для отправки сообщений.

    Возвращает:
    - None
    """
    await callback.message.edit_reply_markup()
    await callback.message.delete()
    await callback.message.answer(text='Подписки сохранены!')
    apscheduler.add_job(send_message_time,
                        trigger='cron',
                        hour=16,
                        kwargs={'bot': bot, 'session': session})


@user_private_router.callback_query(F.data.startswith('subscribe_'))
async def subscribe_topic(callback: CallbackQuery, session: AsyncSession):
    """
    Обработчик нажатия на кнопку с категорией. Убирает inline-клавиатуру с выбором категорий
    и сохраняет выбранную категорию в базу данных.

    Аргументы:
    - callback (CallbackQuery): Объект сообщения от пользователя.
    - session (AsyncSession): Сессия асинхронного соединения с базой данных.

    Возвращает:
    - None
    """

    topic = callback.data.split("_")[-1]

    curr = await db_query.orm_get_id_category_by_name(session, topic)
    await db_query.orm_add_users_category(session, str(callback.from_user.id), curr)


@user_private_router.callback_query(F.data.startswith('restart_no'))
async def restart_no(callback: CallbackQuery):
    """
    Обработчик нажатия на кнопку 'Нет' при рестаре выбора категорий. Удаляет
    выведенную клавиатуру и подпись к ней.

    Аргументы:
    - callback (CallbackQuery): Объект сообщения от пользователя.

    Возвращает:
    - None
    """
    await callback.message.edit_reply_markup()
    await callback.message.delete()


@user_private_router.callback_query(F.data.startswith('restart_yes'))
async def restart_yes(callback: CallbackQuery, session: AsyncSession):
    """
    Обработчик нажатия на кнопку 'Да' при рестаре выбора категорий. Удаляет
    выведенную клавиатуру и подпись к ней, удаляет старые записи о категориях
    данного пользователя и вносит новые в базу данных.

    Аргументы:
    - callback (CallbackQuery): Объект сообщения от пользователя.
    - session (AsyncSession): Сессия асинхронного соединения с базой данных.

    Возвращает:
    - None
    """
    await callback.message.edit_reply_markup()
    await db_query.orm_delete_users_category(session, str(callback.from_user.id))
    await callback.message.answer("Все подписки удалены!")
    await callback.message.answer("Выбери категории для отслеживания:",
                                  reply_markup=reply.categories_kb)


async def send_message_time(bot: Bot, session: AsyncSession) -> None:
    """
    Генерирует и рассылает дайджесты последних новостей по категориям.

    Аргументы:
    - bot (Bot): Объект бота для рассылки сообщений пользователям.
    - session (AsyncSession): Сессия асинхронного соединения с базой данных.

    Возвращает:
    - None
    """
    users_chat_id = await db_query.orm_get_list_of_users(session)
    latest_news = await db_query.orm_get_latest_news_by_categories(session)
    category = await db_query.orm_get_list_of_category(session)

    for user in users_chat_id:
        message = ""
        subscriptions = await db_query.orm_get_users_categories(session, str(user))
        for subscription in subscriptions:
            message += "\n*'{}'* \n".format(category[int(subscription)])
            for news in latest_news[int(subscription)]:
                message += f"{news.title}: {news.url}\n"
        message += "\n Все последние новости и краткую информацию по ним ты можешь найти в файле под этим сообщением."

        await send_docx_to_subscribers(session, bot)

        await bot.send_message(chat_id=user, text=message, parse_mode=ParseMode.MARKDOWN)


async def send_docx_to_subscribers(session: AsyncSession, bot: Bot) -> None:
    """
    Составляет дайджест новостей в формате документа docx и рассылает пользователям.

    Аргументы:
    - bot (Bot): Объект бота для рассылки документов пользователям.
    - session (AsyncSession): Сессия асинхронного соединения с базой данных.

    Возвращает:
    - None
    """
    users_chat_id = await db_query.orm_get_list_of_users(session)

    latest_news = await db_query.orm_get_latest_news_by_categories(session)
    category = await db_query.orm_get_list_of_category(session)

    today = datetime.date.today().isoformat()

    for user in users_chat_id:
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc = docx.Document()

            # колонтикулы
            header = doc.sections[0].header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = f"Дайджест Финтех новостей на {today}"
            run = header_para.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(11)

            subscriptions = await db_query.orm_get_users_categories(session, str(user))

            for subscription in subscriptions:
                # название категории
                category_paragraph = doc.add_paragraph()
                category_run = category_paragraph.add_run(category[int(subscription)])
                category_run.bold = True
                category_run.font.size = Pt(20)
                category_run.font.name = 'Arial'
                category_run.font.color.rgb = RGBColor(0, 0, 0)
                category_run.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for news in latest_news[int(subscription)]:

                    # форматирование для новостного заголовка
                    doc.add_paragraph()
                    new_title = doc.add_paragraph()
                    new_title_run = new_title.add_run(news.title + ":")
                    new_title_run.bold = True
                    new_title_run.font.name = 'Arial'
                    new_title_run.font.size = Pt(11)
                    new_title_run.font.color.rgb = RGBColor(0, 136, 238)

                    # форматирование для текста саммари новости
                    if news.summary is not None:
                        new_content = doc.add_paragraph()
                        new_content_run = new_content.add_run(news.summary)
                        new_content_run.font.size = Pt(11)
                        new_content_run.font.color.rgb = RGBColor(0, 0, 0)
                        new_content_run.font.name = 'Arial'

                    # форматирование для ссылка на исходник
                    new_url = doc.add_paragraph()
                    new_url_run = new_url.add_run("Ссылка на источник: " + news.url)
                    new_url_run.font.name = 'Arial'
                    new_url_run.font.italic = True
                    new_url_run.font.size = Pt(11)
                    new_url_run.font.color.rgb = RGBColor(0, 0, 0)

            doc.save(temp_file.name)

            with open(temp_file.name, mode='rb') as f:
                text_file = FSInputFile(f.name, filename=f'news_digest_{today}.docx')
                await bot.send_document(user, text_file)
